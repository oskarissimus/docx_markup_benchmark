from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from typing import Iterable

from docx import Document  # type: ignore[import-not-found]
from pdf2image import convert_from_path  # type: ignore[import-not-found]


def new_doc() -> Document:
    return Document()


def add_table(doc: Document, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    # Add borders via built-in style
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    # Ensure at least one paragraph per cell
    for row in table.rows:
        for cell in row.cells:
            if not cell.paragraphs:
                cell.add_paragraph("")
    return table


def set_cell_text(cell, text: str) -> None:
    cell.text = text


def add_runs(cell, parts: Iterable[str]) -> None:
    # Clear existing content, then add runs to the first paragraph
    cell.text = ""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
    for part in parts:
        para.add_run(part)


def merge(table, rs: int, cs: int, re: int, ce: int) -> None:
    table.cell(rs, cs).merge(table.cell(re, ce))


def save(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def export_artifacts(src_docx: Path, out_dir: Path, stem: str) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    # Always copy DOCX
    target_docx = out_dir / f"{stem}.docx"
    if src_docx.resolve() != target_docx.resolve():
        target_docx.write_bytes(src_docx.read_bytes())

    # Try soffice pdf conversion
    soffice = shutil.which("soffice")
    if not soffice:
        return
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(src_docx),
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        pdf_path = out_dir / (src_docx.stem + ".pdf")
        if pdf_path.exists():
            images = convert_from_path(str(pdf_path))
            for idx, im in enumerate(images, start=1):
                im.save(out_dir / f"{stem}_{idx}.png")
    except Exception:
        # Best-effort; ignore failures
        pass


